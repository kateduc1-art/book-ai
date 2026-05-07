"""
Export service: generates Markdown and DOCX files from a Chapter.
"""
from __future__ import annotations

import re
from pathlib import Path

from sqlalchemy.orm import Session

from app.core.config import settings
from app.core.logging import get_logger
from app.models.chapter import Chapter

logger = get_logger(__name__)


def _ensure_export_dir() -> Path:
    export_dir = Path(settings.export_dir)
    export_dir.mkdir(parents=True, exist_ok=True)
    return export_dir


def export_markdown(db: Session, chapter: Chapter) -> str:
    """Write chapter content to a .md file and return the file path."""
    export_dir = _ensure_export_dir()
    safe_title = re.sub(r"[^\w\-]", "_", chapter.title)[:60]
    filename = f"{chapter.id}_{safe_title}.md"
    file_path = export_dir / filename

    content = chapter.rewritten_markdown or chapter.markdown_content or ""
    file_path.write_text(content, encoding="utf-8")

    chapter.export_markdown_path = str(file_path)
    db.commit()

    logger.info(f"Exported Markdown: {file_path}")
    return str(file_path)


def export_docx(db: Session, chapter: Chapter) -> str:
    """Write chapter content to a .docx file and return the file path."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    export_dir = _ensure_export_dir()
    safe_title = re.sub(r"[^\w\-]", "_", chapter.title)[:60]
    filename = f"{chapter.id}_{safe_title}.docx"
    file_path = export_dir / filename

    content = chapter.rewritten_markdown or chapter.markdown_content or ""
    lines = content.splitlines()

    doc = DocxDocument()

    # Document title
    heading = doc.add_heading(chapter.title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line in lines:
        stripped = line.strip()

        if not stripped or stripped == "---":
            doc.add_paragraph("")
            continue

        # Skip HTML comments
        if stripped.startswith("<!--"):
            continue

        # Markdown headings
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        # Blockquote
        elif stripped.startswith("> "):
            p = doc.add_paragraph(stripped[2:])
            p.style = "Quote"
        # Bullet list
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        # Page reference footnote
        elif stripped.startswith("*[p.") and stripped.endswith("]*"):
            p = doc.add_paragraph(stripped[1:-1])
            run = p.runs[0]
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        # Normal paragraph
        else:
            doc.add_paragraph(stripped)

    doc.save(str(file_path))

    chapter.export_docx_path = str(file_path)
    db.commit()

    logger.info(f"Exported DOCX: {file_path}")
    return str(file_path)
